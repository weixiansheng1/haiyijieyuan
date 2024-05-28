# -*- coding: utf-8 -*-
"""
Created on Sat May 20 15:40:26 2023

@author: 11298
"""

import os
import pandas
import time
import pandas as pd
import re
import string
from datetime import datetime, timedelta
from docxtpl import DocxTemplate #pip install docxtpl
import docx
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.table import _Cell
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches
#获取当前日期
now = time.strftime("%Y-%m-%d",time.localtime(time.time()))
now1 = time.strftime("%Y年 %m月 %d日",time.localtime(time.time()))
now2 = time.strftime("%Y0%m",time.localtime(time.time()))
now3 = time.strftime("%Y年%m月份",time.localtime(time.time()))
now4 = time.strftime("%Y年%m月",time.localtime(time.time())).replace('年0','年')
now5 = time.strftime("%Y年%m月%d日",time.localtime(time.time())).replace('年0','年').replace('月0','月')
# =============================================================================
#                      整改通知书
# =============================================================================
#读取word模板
zhenggai = DocxTemplate(r"E:\海宜洁源安全共享文件\洁源公司安全生产档案资料\8.隐患排查与治理\隐患台账自动化输入\【模板】-勿动\海宜洁源公司安全检查整改通知书（餐厨垃圾处理一期项目）.docx")

# 替换整改通知书编号,只能在txt文件修改编号
folder_path = 'E:/海宜洁源安全共享文件/洁源公司安全生产档案资料/8.隐患排查与治理/隐患台账自动化输入/【生成】-勿动'
file_names = [f for f in os.listdir(folder_path) if f.startswith('上次生成整改') and f.endswith('.txt')] #查看文件夹下的txt文件，取其文件名给编号

#读取隐患内容输入，生成dataframe
yinhuanneirong = pd.read_excel('E:\海宜洁源安全共享文件\洁源公司安全生产档案资料\8.隐患排查与治理\隐患台账自动化输入\隐患内容输入.xlsx')

#读取单元格所对应的dataframe，将其对应其相应的数据，例：检查单位单元格（1，1）其对应的行是0，分隔符"\n"，则调用函数danyuange(1,1,0,'\n')
def zifuchuan_duiying(zifu,duiying_row,fen_ge_fu):
    m_list = re.findall(r'\w|\d', zifu)#分割字符串，生成list
    m_dict = {'1': 1, '2': 2, '3': 3, '4': 4, '5': 5, '6': 6,'7': 7, '8': 8, '9': 9, 'a': 10, 'b': 11, 'c': 12,'d': 13,'e': 14, 'f': 15, 'g': 16,'h': 17, 'i': 18,'j': 19,'k': 20,     }
    danyuange_list = [m_dict[item] for item in m_list]#以list对应dict key——》value
    duiying_neirong = yinhuanneirong.iloc[duiying_row,danyuange_list].tolist()#对应单元格后转list
    table = str.maketrans('', '', string.digits+ string.ascii_lowercase + string.ascii_uppercase)#删除数字、字母
    my_string = fen_ge_fu.join(str(item).translate(table) for item in duiying_neirong)
    return my_string

def danyuange(row,column,duiying_row,fen_ge_fu):#（输入行，列，对应的行）
    du_danyuange = yinhuanneirong.iloc[row,column]#通过iloc使用索引调用
    danyuange_num = str(du_danyuange)
    c = zifuchuan_duiying( danyuange_num, duiying_row, fen_ge_fu)
    return c

#对隐患内容的dataframe切片并重新索引
nan_suoyin = yinhuanneirong.iloc[13:,1].isna().idxmax()#判断非nan,需要在excel表格的隐患内容增加一个空格表示nan
yinhuan = yinhuanneirong.iloc[12:nan_suoyin,1:7].reset_index(drop=True)

#隐患内容拼接
    #判断责任单位
def check_letter(val):#字符串是否包含x字母
    if 'x' in val:
        a = val.replace('x','') #删x
        a = str(a)
        a = '责任单位：' +  zifuchuan_duiying(a, 0, '、')
        return  a
    elif 'r' in val:#对应责任人
        b  = val.replace('r','责任人：')
        return b
    elif 'm' in val:#对应责任部门
        d = val.replace('m','责任部门：')
        return d
    else:
        val = str(val)
        c = zifuchuan_duiying(val, 0, '、')
        return c

#如果单元格内容为1，则为即时整改，如果是日期转化数字，则将excel日期转换为字符串
def riqizhuanhuan(d):
    if isinstance(d, datetime):
        data_zhuanhuan = d.strftime('%Y年%m月%d日前').replace('年0','年').replace('月0','月')
    elif d == 1:
        data_zhuanhuan = '即时整改'
    else:
        date = datetime(1900, 1, 1) + timedelta(days=d)
        data_zhuanhuan = date.strftime('%Y年%m月%d日前').replace('年0','年').replace('月0','月')
    return data_zhuanhuan

#存在问题、整改要求，合并相应的字符 
quanbuyinhuan = ''
quanbuzhenggaiyaoqiu = ''
for ij in range(1,len(yinhuan.index)):
    #存在问题
    panduan_x = check_letter('{}'.format(yinhuanneirong.iloc[12+ij,3]))
    neirong = str(ij)  +"." + yinhuan.iloc[ij,0] + str("({})".format(panduan_x))
    quanbuyinhuan = quanbuyinhuan + neirong + '\n'
    #整改要求
    zhenggaiyaoqiu = str(ij) + '.' +yinhuan.iloc[ij,3] +str('(整改期限：{})'.format(riqizhuanhuan(yinhuan.iloc[ij,4])))
    quanbuzhenggaiyaoqiu = quanbuzhenggaiyaoqiu + zhenggaiyaoqiu +'\n'
    #问题描述
    wentimiaoshu = yinhuan.iloc[ij,0]
    pass

#整改通知书附件内容
#参考文章https://zhuanlan.zhihu.com/p/366902690
dict2 = yinhuan.iloc[1:,0].to_dict()
fujianneirong_fenge = [{'问题描述': dict2[ki],'序号': str(ki),'图片':'' }
                     for ki ,a in dict2.items()]#ki:键  a ：值

#替换函数
context = {"编号":"{}{}".format(file_names[0][12:-6],int(file_names[0][19:-4])+1),#【替换编号】将数字编号拆分，转换成int后每次加1
           "检查时间": now5,
           '被检查单位': '{}'.format(danyuange(1,1,0,'\n')),#需要使用format直接变成字符串才能跟随文档格式！！！
           '检查人员': '{}'.format(danyuange(4,1,3,'、')),
           '检查地点':'{}'.format(danyuange(7,1,6,'\n')) ,
           '其他要求':'{}'.format(danyuange(10,1,9,'')) ,
           '存在问题': '{}'.format(quanbuyinhuan),
           '整改要求': '{}'.format(quanbuzhenggaiyaoqiu),
           '附件':fujianneirong_fenge,#整改通知书附件,
           }

zhenggai.render(context)#渲染

#读取图片地址
source_tupian = r"E:/海宜洁源安全共享文件/洁源公司安全生产档案资料/8.隐患排查与治理/隐患台账自动化输入"
file_tupian = r'E:\\海宜洁源安全共享文件\\洁源公司安全生产档案资料\\8.隐患排查与治理\\隐患台账自动化输入'
img_files = os.listdir(source_tupian)
img_paths = {}
for file in img_files:
    if file.endswith(('.png', '.jpg', '.jpeg')):#endswith()方法只能接受一个字符串作为参数
        path = os.path.join(source_tupian, file)
        num = file.split('.')[0]  # 获取文件名中的数字
        img_paths[num] = file_tupian + path[46:]


# 定义一个函数，根据单元格内容返回对应的图片路径
def get_img_path(cell_value):
    result = []
    if isinstance(cell_value, int):
        cell_value = str(cell_value)
        result.append(img_paths.get(cell_value))
    else:
        cell_value = cell_value.split('、')
        for item in cell_value:
            result.append(img_paths.get(item))
    return result

#adfdasf = get_img_path(yinhuan.iloc[1,1])#测试图片

#表格读取
table = zhenggai.tables[1]# 获取第biaoge个表格
print (table)
#插入图片至单元格
xuhao =1
dict3 = yinhuan.iloc[1:,1]
for i in range(1,len(yinhuan)):
    a_s = get_img_path(yinhuan.iloc[i,1])
    for i in range(len(a_s)):
        a = a_s[i]
        cell = table.cell(xuhao,2)# 在第二个表格的第一行第一列中插入图片
        cell.paragraphs[0].add_run().add_picture('{}'.format(a), width=docx.shared.Cm(10))
        xuhao = xuhao + 1 

'''字体大小对应pt
初号44pt
小初36pt
号26pt
小一24pt
二号22pt
小二18pt
三号16pt
小三15pt
四号14pt
小四12pt
五号10.5pt
八号5
七号5.5
字号‘小六’对应磅值6.5
字号‘六号’对应磅值7.5
字号‘小五’对应磅值9
'''
# 统计段落数和行数
def tongji_d_h(wenjian,mingcheng):
    num_paragraphs = len(wenjian.paragraphs)
    num_lines = 0
    for paragraph in wenjian.paragraphs:
        num_lines += len(paragraph.text.split('\n'))
        # 统计表格数
    num_tables = len(wenjian.tables)
    for table in wenjian.tables:    
        row_count = len(table.rows)
        column_count = len(table.columns)
        print(f'{mingcheng}模板的第{table}个表格中有{row_count}行，{column_count}列')
    print(f"{mingcheng}模板中有 {num_paragraphs} 个段落和 {num_lines} 行文本, {num_tables} 个表格。")



# =============================================================================
#                 隐患整改确认单
# =============================================================================

#==== 设置 table 的边框，用法与 cell 类似
def set_table_boarder(table, **kwargs):
    """
    Set table`s border
    Usage:
    set_table_border(
        cell,
        top={"sz": 12, "val": "single", "color": "#FF0000"},
        bottom={"sz": 12, "color": "#00FF00", "val": "single"},
        left={"sz": 24, "val": "dashed"},
        right={"sz": 12, "val": "dashed"},
    )
    """
    borders = OxmlElement('w:tblBorders')
    for tag in ('bottom', 'top', 'left', 'right', 'insideV', 'insideH'):
        edge_data = kwargs.get(tag)
        if edge_data:
            any_border = OxmlElement(f'w:{tag}')
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    any_border.set(qn(f'w:{key}'), str(edge_data[key]))
            borders.append(any_border)
            table._tbl.tblPr.append(borders)
    # 将table 的所有单元格四个边设置为 0.5 镑, 黑色, 实线
def set_table_singleBoard(table): return set_table_boarder(
    table,
    top={"sz": 4, "val": "single", "color": "#000000"},
    bottom={"sz": 4, "val": "single", "color": "#000000"},
    left={"sz": 4, "val": "single", "color": "#000000"},
    right={"sz": 4, "val": "single", "color": "#000000"},
    insideV={"sz": 4, "val": "single", "color": "#000000"},
    insideH={"sz": 4, "val": "single", "color":  "#000000"}
)
#============
#判断生成海宜还是洁源
asfffff = yinhuanneirong.iloc[11,1]
if asfffff == 1 :
    shengcheng_1 = '整改确认单-自查.docx'
elif asfffff == 3:
    shengcheng_1 = '整改确认单-自查.docx'
elif asfffff == 2 :
    shengcheng_1 = '整改确认单-海宜查.docx'
elif asfffff == 4:
    shengcheng_1 = '整改确认单-自查.docx'
else:
    shengcheng_1 = '整改确认单-自查.docx'
    pass


#判断是否有文件存在，如果不存在：从模板文件生成；如果存在:判断本月是否超过确认单月，补充【生成】文件
querendan_file_names = [f for f in os.listdir(folder_path) if f.endswith('{}'.format(shengcheng_1))]  # 获取所有以 整改确认单.docx结尾的文件名列表
if querendan_file_names != []:#如果存在:补充【生成】文件
    file_path = os.path.join(r'E:/海宜洁源安全共享文件/洁源公司安全生产档案资料/8.隐患排查与治理/隐患台账自动化输入/【生成】-勿动/{}'.format(querendan_file_names[0]))
    querendan = Document(file_path)
    qr_table = querendan.tables[0]
    num_rows = len(qr_table.rows)-4
    #增加行，并填入内容
    for i in range(1,len(yinhuan)):
        num_rows = num_rows+1
        new_row = qr_table.add_row()
        new_row.cells[0].text = '{}'.format(num_rows)
        new_row.cells[1].text = yinhuan.iloc[i,0]
        new_row.cells[2].text = now5
        new_row.cells[3].text = yinhuan.iloc[i,3]
        # 设置新行单元格的格式
        for cell in new_row.cells:
       # 设置单元格格式
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.paragraphs[0].style.font.name = 'FangSong_GB2312'
            cell.paragraphs[0].style.font.size = Pt(14)
        #添加图片
            a_s = get_img_path(yinhuan.iloc[i,1])
        for i in range(len(a_s)):
            a = a_s[i]
            new_row.cells[4].paragraphs[0].add_run().add_picture('{}'.format(a), width=docx.shared.Cm(6))
        qr_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_singleBoard(qr_table)#画框线
    print("{querendan_file_names} 文件已存在，已补充完毕！")
else:# 如果不存在：从模板文件生成
    querendan  = DocxTemplate(r'E:\海宜洁源安全共享文件\洁源公司安全生产档案资料\8.隐患排查与治理\隐患台账自动化输入\【模板】-勿动\洁源公司整改确认单.docx')
    #确认单
    querendan_zhenggaicuoshi = yinhuan.iloc[1:,3].to_dict()
    querendan_zhenggaicuoshi = {str(k): v for k, v in querendan_zhenggaicuoshi.items()}#字典键转为字符串
    querendan_sc = [{'问题描述': dict2[ki],'序号': str(ki),'图片':'' ,'整改措施':'','检查时间':now5}
                         for ki ,a in dict2.items()]#ki:键  a ：值
    #添加querendan_sc中“整改措施”键的值
    for item in querendan_sc:
        key = item['序号']
        if key in querendan_zhenggaicuoshi:
            item['整改措施'] = querendan_zhenggaicuoshi[key]
    context2 = {'确认单编号':now2,
                '确认单': querendan_sc,
                }
    querendan.render(context2)
    #图片生成
    qrd_table = querendan.tables[0]# 获取第biaoge个表格
    #插入图片至单元格
    for i in range(1,len(yinhuan)):
        a_s = get_img_path(yinhuan.iloc[i,1])
        for i in range(len(a_s)):
            a = a_s[i]
            cell = qrd_table.cell(i+4,4)# 在表格的第5行第4列中插入图片
            cell.paragraphs[0].add_run().add_picture('{}'.format(a), width=docx.shared.Cm(6))

    print("整改确认单文件不存在，从模板文件生成。")


# =============================================================================
#                               检查通报
# =============================================================================
output_dir = r"E:\海宜洁源安全共享文件\洁源公司安全生产档案资料\8.隐患排查与治理\隐患台账自动化输入\【生成】-勿动"
output_path = os.path.join(output_dir, "关于珠海市海宜洁源餐厨垃圾处置有限公司{}安全生产专项工作检查情况的通报.docx ".format(now4))
abc = 0
#判断是否已经生成检查通报
jiancha_file_names = [f for f in os.listdir(folder_path) if f.endswith('月安全生产专项工作检查情况的通报.docx')]  # 获取所有以 整改确认单.docx结尾的文件名列表
if jiancha_file_names == []:#如果不存在
    #读取word模板
    tongbao1 = docx.Document(r"E:\海宜洁源安全共享文件\洁源公司安全生产档案资料\8.隐患排查与治理\隐患台账自动化输入\【模板】-勿动\安全生产专项工作检查情况的通报.docx")
    #设置检查序号
    tongbao1.save(output_path)
    pass


tongbao2 =  docx.Document(output_path)
# 统计指定段落中的特定字符串数量
def count_target_text(start_text, end_text, target_text):
    count = 0
    start_found = False
    for paragraph in tongbao2.paragraphs:
        if start_text in paragraph.text:
            start_found = True
        elif end_text in paragraph.text:
            break
        elif start_found:
            if target_text in paragraph.text:
                count += 1
    return count


if asfffff == 1 :#洁源
    chazhao = '收运、调度、车辆安全检查情况'
    # 调用函数统计特定字符串数量
    count = count_target_text('各停车场、洗车场、收运点安全检查情况', '收运、调度、车辆安全检查情况','检查时间')
    abc = count + 1
elif asfffff == 3:#收运
    chazhao = '海宜公司检查情况'
    count = count_target_text('收运、调度、车辆安全检查情况', '海宜公司检查情况','检查时间')
    abc = count + 1
elif asfffff == 2 :#海宜
    chazhao = '食品安全检查情况'
    count = count_target_text('海宜公司检查情况','食品安全检查情况','检查时间')
    abc = count + 1
elif asfffff == 4:#食品
    chazhao = '珠海市城管局检查情况'
    count = count_target_text('食品安全检查情况','珠海市城管局检查情况','检查时间')
    abc = count + 1
elif asfffff == 5:##城管
    chazhao = '富山应急局检查情况'
    count = count_target_text('珠海市城管局检查情况','富山应急局检查情况','检查时间')
    abc = count + 1
elif asfffff == 6:##富山
    chazhao = '           '
    count = count_target_text('富山应急局检查情况','           ','检查时间')
    abc = count + 1
else:
    pass


tongbao3 =  docx.Document(output_path)
# 查找指定的文本
for paragraph in tongbao3.paragraphs:
    if chazhao in paragraph.text:
        # 在指定文本上方插入新段落
        new_paragraph = paragraph.insert_paragraph_before('')
        # 设置新段落格式
        new_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        new_paragraph.paragraph_format.first_line_indent = Pt(28)#首行缩进
        # 添加检查时间和存在问题的文本
        run = new_paragraph.add_run('{a}.检查时间：{b}'.format(a = abc,b ='{{检查时间}}'))
        run.font.name = u'仿宋'
        run.font.size = Pt(16)
        run.font.bold = True
        # 设置中文字体
        # 需导入 qn 模块from docx.oxml.ns import qn
        run.font.element.rPr.rFonts.set(qn('w:eastAsia'),'仿宋')#https://www.jianshu.com/p/8f15e3f2f9e6
        #==========================================================
        new_paragraph = paragraph.insert_paragraph_before('')
        # 设置新段落格式
        new_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        new_paragraph.paragraph_format.first_line_indent = Pt(28)#首行缩进
        run1 = new_paragraph.add_run('存在问题:')
        run1.font.name = u'仿宋'
        run1.font.size = Pt(16)
        run1.font.element.rPr.rFonts.set(qn('w:eastAsia'),'仿宋')#https://www.jianshu.com/p/8f15e3f2f9e6
        #==========================================================
        new_paragraph = paragraph.insert_paragraph_before('')
        # 设置新段落格式
        new_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        new_paragraph.paragraph_format.first_line_indent = Pt(28)#首行缩进
        run1 = new_paragraph.add_run('{{厂存问及整改}}')
        run1.font.name = u'仿宋'
        run1.font.size = Pt(16)
        run1.font.element.rPr.rFonts.set(qn('w:eastAsia'),'仿宋')#https://www.jianshu.com/p/8f15e3f2f9e6
        break
#tongbao3.save(output_path)

tongbao4 = DocxTemplate(output_path)
#本月检查情况及整改意见
cang_neirong = ''
for ij in range(1,len(yinhuan.index)):
    #存在问题
    panduan_x = check_letter('{}'.format(yinhuanneirong.iloc[12+ij,3]))
    cang_cunzaiwenti = '（{}）'.format(ij) + yinhuan.iloc[ij,0] + str("({})".format(panduan_x))+'\n'
    #整改要求
    zhenggaiyaoqiu = '    整改意见：' + yinhuan.iloc[ij,3] + str('(整改期限：{})'.format(riqizhuanhuan(yinhuan.iloc[ij,4])))
    #合并存在问题、整改要求
    hebing_cang = cang_cunzaiwenti + zhenggaiyaoqiu
    cang_neirong = cang_neirong + hebing_cang +'\n'
    pass


#替换函数
context = {
           '月份': now4,
           "检查时间": now5,
           '厂存问及整改': '{}'.format(cang_neirong),#需要使用format直接变成字符串才能跟随文档格式！！！
           }

tongbao4.render(context)#渲染











# =============================================================================
#           保存生成内容
# =============================================================================
    
tongji_d_h(zhenggai, '整改通知书')
tongji_d_h(querendan, '整改确认单')

#保存到【输出】文件夹，并加入今天的日期
output_dir = r"E:\海宜洁源安全共享文件\洁源公司安全生产档案资料\8.隐患排查与治理\隐患台账自动化输入\【生成】-勿动"
output_path = os.path.join(output_dir, "海宜洁源公司安全检查整改通知书（餐厨垃圾处理一期项目）{}.docx ".format(now))
zhenggai.save(output_path)
output_path1 = os.path.join(output_dir, "海宜洁源公司{}{}".format(now3,shengcheng_1))
querendan.save(output_path1)
output_path = os.path.join(output_dir, "关于珠海市海宜洁源餐厨垃圾处置有限公司{}安全生产专项工作检查情况的通报.docx ".format(now4))
tongbao4.save(output_path)
# =============================================================================
# 移动生成好的文件至相应的文件夹，并生成.txt记录
# =============================================================================

#查看某文件夹下是否存在按月份生成的文件夹，如果没有则生成
def wenjianjia(wenjianjia_path):
    nowmonth = datetime.datetime.now()
    current_month = nowmonth.strftime("%Y年-%m月")
    folder_path = os.path.join(wenjianjia_path, current_month) # 构建文件夹路径
    if not os.path.exists(folder_path):# 判断文件夹是否存在，如果不存在则创建它
        os.makedirs(folder_path)
    return folder_path

#从【输出】文件夹剪切某文件至另一文件夹
def jianqie (a,d):
    source_folder = "D:/8.隐患排查与治理/隐患台账自动化输入/【生成】-勿动"# 源文件夹路径
    destination_folder = wenjianjia(a) # 调用wenjianjia函数，a是目标文件夹路径
    filename = d # 要移动的文件名
        # 构建源文件路径和目标文件路径
    source_path = os.path.join(source_folder, filename)
    destination_path = os.path.join(destination_folder, filename)
        # 移动文件
    os.rename(source_path, destination_path)


#移动文件的目标文件夹，月份已自动匹配
#mubiaowenjian = "D:\8.隐患排查与治理\8.4整改通知书" #仅填两个层级文件夹
#目标文件
#jianqiewenjian = "海宜洁源公司安全检查整改通知书（餐厨垃圾处理一期项目）{}.docx ".format(now)
#mubiaowenjian = jianqie(mubiaowenjian,jianqiewenjian) #移动


#删除上次，生成下次txt文件记录上一次生成(编号不能高或低于2位数)
file_path = 'E:/海宜洁源安全共享文件/洁源公司安全生产档案资料/8.隐患排查与治理/隐患台账自动化输入/【生成】-勿动/{}'.format(file_names[0])
if os.path.exists(file_path):# 如果文件存在，则删除它
    os.remove(file_path)
    print(f"{file_path} 文件已成功删除！")
else:
    print(f"{file_path} 文件不存在。")
    
with open('E:/海宜洁源安全共享文件/洁源公司安全生产档案资料/8.隐患排查与治理/隐患台账自动化输入/【生成】-勿动/上次生成整改通知书编号：{}{}.txt'.format(file_names[0][12:-6],int(file_names[0][19:-4])+1), 'w') as f:
# 写入文本内容
    f.write('Hello World!')


