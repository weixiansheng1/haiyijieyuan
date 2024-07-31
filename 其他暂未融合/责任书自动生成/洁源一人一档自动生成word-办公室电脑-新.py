# -*- coding: utf-8 -*-
"""
Created on Mon Apr 24 12:17:01 2023 by廖振威

"""

# -*- coding: utf-8 -*-
"""
使用说明：

本功能用于合并并排序excel表格，并且美化excel
1.将钉钉上下载的培训记录和考试的excel表格放入同一个文件夹（所有的文件名称都要删除跟在的后面杂七杂八的字符,包括最后的“-”）
2.将下载好的培训记录excel文件名的第一个数字确定为该课程的学时，考试记录不用填
3.如果模板路径有变动，则要复制路径到提示错误的位置
4.d = x是月份，改成x月份
5.运行
6.程序生成该月的课程的完成情况excel表格
7.程序生成一份以部门、姓名字母排序好的word一人一档

注意：
1.现场培训需自己手动在生成好后的word内添加
2.培训记录excel表要与考试记录excel表相匹配，不然会报错：list index out of range
3.如果出现多个人有姓名的前两个字相同，则在代码的替换中填入替换字才能捋顺顺序

"""


import os
import pandas as pd
import openpyxl 
import pypinyin
from   openpyxl                 import Workbook 
from   pypinyin                 import lazy_pinyin
from   docx                     import Document                 # docx操作库,需要安装 pip install python-docx,老版本报错，卸载原来安装的docx
from   docx.shared              import Cm                       # 引入厘米库
from   docx.shared              import Pt                       # 引入字体大小调整
from   docx.oxml.ns             import qn                       # 字体库
from   docx.enum.text           import WD_PARAGRAPH_ALIGNMENT   # 段落对齐
from   docx.enum.table          import WD_TABLE_ALIGNMENT
from   docx.shared              import RGBColor
from   docx.enum.text           import WD_LINE_SPACING          # 设置特殊行间距
from   openpyxl.styles          import Font,PatternFill, Border, Side
from   openpyxl.utils.dataframe import dataframe_to_rows


# =============================================================================
#  ============== 新生成d月份一人一档统计表格 ===================================
# =============================================================================
d = 6 #设置月份

os.chdir(r"E:\海宜洁源安全共享文件\标准化自评\【5】教育培训\5.1教育培训管理\3-培训记录及评估\1-日常培训记录\一人一档内容编制\{}月份".format(d))# 定义钉钉下载的源数据路径，该路径下不得有非xlsx文件

wjm  = os.listdir()                        # 读取文件下的所有文件名称
xylb =  [s for s in wjm if "学员统计" in s] # 读取学员统计列表
kstj =  [s for s in wjm if "考试统计" in s] # 读取考试统计列表
xylb.sort()                                # xylb排序对齐kstj，优化匹配
kstj.sort()                                # xylb排序对齐kstj

# ==================钉钉培训记录表生成==========================================

o = ['A','B','C','D','E']                                                                   # 新生成的钉钉教育培训记录表的列
aligna = openpyxl.styles.Alignment(horizontal= 'center',vertical='center',wrap_text = True) # 自动换行，居中
border = Border(left=Side  (border_style='thin', color='000000'),                           # 单元格cell的边框设置
                right=Side (border_style='thin', color='000000'),
                top=Side   (border_style='thin', color='000000'),
                bottom=Side(border_style='thin', color='000000'))

# ========合并os.chdir路径下的多个excel表格至一个表==============================
i = 0
v = 0                    #考试统计控制
xueyuan = 0              #学员列表控制
hebing = pd.DataFrame()  #创建空DataFrame
for excel_name in xylb:  #遍历读取的文件名

    c   = pd.read_excel(excel_name,sheet_name="第一页",usecols=[0,2,4,6,14,16])       # 读取课程学习记录的excel，使 用0,2,4,6，13列的值
    
    #如果当月完全没有考试，则用空值替换，如果有考试则读取考试统计的excel
    print('运行到第{}个excel，共有{}个excel'.format(v,len(xylb)))

    if kstj == []:
            ksv = pd.read_excel(excel_name,sheet_name="第一页",usecols=[0,17])         # 为并列对齐，用课程学习记录中的0列和17列，做 姓名+空成绩替代无考试成绩excel情况,还需更新下面的‘不打印’
            ksv['最新一次成绩'] = pd.Series([float('NaN')for x in range(len(c))])
    
    else:
        
        for kaoshi_name in kstj:       # 循环匹配考试的excel名字与学员统计的excel名                                                            
            if excel_name[1:-11] == kaoshi_name[0:-11] : 
                ksv = pd.read_excel(kaoshi_name,sheet_name="第一页", usecols=[0,7])  # [学院姓名，最新成绩]
                print('当前考试：{}'.format(kaoshi_name[0:-11]))
                print('当前学员：{}'.format(excel_name[1:-11]))
                break
            else:
                ksv = pd.read_excel(excel_name,sheet_name="第一页",usecols=[0,17])    # 为并列对齐，用课程学习记录中的0列和19列，做 姓名+空成绩替代无考试成绩excel情况
                ksv['最新一次成绩'] = pd.Series([float('NaN')for x in range(len(c))])
                

    g   = c
    # ============生成excel总表，供word使用===============
    g['指派/加入时间'] = '{}'.format(g.loc[0,'指派/加入时间'])                      # 取第0行，第'指派/加入时间'列 由于excel表中“指派/加入时间”只是加入时间而非指派时间，故需在dataframe对其列进行修改
    # 新增‘课程名称’列用于datafram内按行添加课程名称
    for i in range(len(g)):                                                       # range(10) 是 0 到 9；len 统计dataframe的index数量，行数
        g.loc[i,'课程名称'] = excel_name                                           # 添加'课程名称'的列
        g.loc[i,'课时']     = excel_name[0]
        
    # 课程合并成绩
    g = pd.merge(g, ksv,how='outer',on='学员姓名')                                 # merge(第一个DataFrame，第二个DataFrame，合并方式，第一个DataFrame合并用的键，第二个DataFrame合并用的键)   
    v = v + 1                                                                     # 用于ksv和excel_name并轨对齐 
    hebing = pd.concat([hebing,g])                                                # 利用concat合并DataFrame
    
    
    # ========美化钉钉培训记录excel======================

    #读取excel中需要的列至dataframe内
    x  = c[['学员姓名','部门','职位','完成情况','完成时间']]
    wb = Workbook()                                # 新建空workbook
    ws = wb.active                                 # 获取活动中的sheet，该为默认

    #将读取的x的dataframe按行填入ws
    for each in dataframe_to_rows(x, index=False, header=True):
        ws.append(each)
    #插入空行在第1行，留待后面合并后填入文件名
    ws.insert_rows(1)

    #按行设置表格的行高为27
    for i in range(1,ws.max_row):
        ws.row_dimensions[i].height = 27
    for i in o:
        for cell in ws[i]:
            cell.alignment = aligna # 表格样式
            cell.border    = border # 表格

    #对第二行加粗
    for t in ws.iter_rows(min_row=2,max_row= 2):
        for cell in t:
            cell.font = Font(bold = True, size = 12, )                                                  # 加粗
            cell.fill = PatternFill(fill_type='darkHorizontal', start_color='99CCFF',end_color='99CCFF')# 单元格颜色

    #设置列宽度
    ws.column_dimensions['A'].width = 18.91
    ws.column_dimensions['B'].width = 31.77
    ws.column_dimensions['C'].width = 10
    ws.column_dimensions['D'].width = 14.13
    ws.column_dimensions['E'].width = 23
    
    #在A1单元格内填入文件名
    ws['A1'] = excel_name[1:-12]
    ws.merge_cells('A1:E1')      # 合并单元格
    wb.save(r"E:\海宜洁源安全共享文件\标准化自评\【5】教育培训\5.1教育培训管理\3-培训记录及评估\1-日常培训记录\{}月份培训\{}" .format(d,excel_name))#储存每一个学员统计表
   
    # ============美化钉钉考试excle ================================
    
lb = ['A','B','C','D','E','F','G','H'] 

for excel_name in kstj:
    ks   = pd.read_excel(excel_name,sheet_name="第一页",usecols=[0,3,4,5,6,7,8,15])         # 读取excel中需要的列至dataframe内
    ksjl = Workbook()                                                                      # 新建空workbook
    hk   = ksjl.active                                                                     # 获取活动中的sheet，该为默认
    
    #将读取的hk的dataframe数据按行填入hk
    for each in dataframe_to_rows(ks, index=False, header=True):
        hk.append(each)
    #插入空行在第1行，留待后面合并后填入文件名
    hk.insert_rows(1)
    for i in range(1,hk.max_row):                # 按行设置表格的行高为28
        hk.row_dimensions[i].height = 28
    for i in lb:
        for cell in hk[i]:
            cell.alignment = aligna              # 表格样式
            cell.border    = border              # 表格
    for t in hk.iter_rows(min_row=2,max_row= 2): # 对第二行加粗
        for cell in t:
            cell.font = Font(bold = True, size = 12, )                                                  # 加粗，12磅
            cell.fill = PatternFill(fill_type='darkHorizontal', start_color='99CCFF',end_color='99CCFF')# 单元格颜色
            
    # 设置列宽度       
    hk.column_dimensions['A'].width = 10
    hk.column_dimensions['B'].width = 13
    hk.column_dimensions['C'].width = 22.25
    hk.column_dimensions['D'].width = 30
    hk.column_dimensions['E'].width = 10
    hk.column_dimensions['F'].width = 8.5
    hk.column_dimensions['G'].width = 8.5
    hk.column_dimensions['H'].width = 19
    # 在A1单元格内填入文件名
    hk['A1'] = '{}-考试情况'.format(excel_name[1:-12]) 
    hk.merge_cells('A1:H1')                                                                                                                                # 合并单元格
    ksjl.save(r"E:\海宜洁源安全共享文件\标准化自评\【5】教育培训\5.1教育培训管理\3-培训记录及评估\1-日常培训记录\{}月份培训\{}" .format(d,excel_name))           # 储存每一个学员统计表
       
    # ============================================================
        
#修改hebing合并后部门的数据值，使工整
#部门替换
hebing.loc[hebing['部门']=='珠海市海宜洁源餐厨垃圾处置有限公司-洁源安环部',['部门']]='珠海市海宜洁源餐厨垃圾处置有限公司-洁源安环部-洁源安环部'
hebing.loc[hebing['部门']=='珠海市海宜洁源餐厨垃圾处置有限公司-洁源收运部',['部门']]='珠海市海宜洁源餐厨垃圾处置有限公司-洁源收运部-洁源收运部'
hebing.loc[hebing['部门']=='珠海市海宜洁源餐厨垃圾处置有限公司-洁源收运部-洁源东区服务点,珠海市海宜洁源餐厨垃圾处置有限公司-洁源收运部',['部门']]='珠海市海宜洁源餐厨垃圾处置有限公司-洁源收运部-洁源东区服务点'
hebing.loc[hebing['部门']=='珠海市海宜洁源餐厨垃圾处置有限公司-调度中心',['部门']]='珠海市海宜洁源餐厨垃圾处置有限公司-洁源调度中-洁源调度中'
hebing.loc[hebing['部门']=='珠海市海宜洁源餐厨垃圾处置有限公司-洁源综合部',['部门']]='珠海市海宜洁源餐厨垃圾处置有限公司-洁源综合部-洁源综合部'
hebing.loc[hebing['部门']=='珠海市海宜洁源餐厨垃圾处置有限公司-洁源收运部-洁源公司押运,珠海市海宜洁源餐厨垃圾处置有限公司-洁源收运部-洁源公司合同,珠海市海宜洁源餐厨垃圾处置有限公司-洁源收运部-洁源东区服务点',['部门']]='珠海市海宜洁源餐厨垃圾处置有限公司-洁源收运部-洁源东区服务点'
hebing.loc[hebing['部门']=='珠海市海宜洁源餐厨垃圾处置有限公司-洁源收运部-洁源公司押运,珠海市海宜洁源餐厨垃圾处置有限公司-洁源收运部-洁源公司合同,珠海市海宜洁源餐厨垃圾处置有限公司-洁源收运部-洁源东区服务点',['部门']]='珠海市海宜洁源餐厨垃圾处置有限公司-洁源收运部-洁源东区服务点'

hebing.loc[hebing['部门']=='珠海市海宜洁源餐厨垃圾处置有限公司-洁源收运部,珠海市海宜洁源餐厨垃圾处置有限公司',['部门']]='珠海市海宜洁源餐厨垃圾处置有限公司-洁源经理室-洁源经理室'
hebing.loc[hebing['部门']=='珠海市海宜洁源餐厨垃圾处置有限公司-洁源生产部,珠海市海宜洁源餐厨垃圾处置有限公司',['部门']]='珠海市海宜洁源餐厨垃圾处置有限公司-洁源经理室-洁源经理室'

hebing.loc[hebing['部门']=='珠海市海宜洁源餐厨垃圾处置有限公司-洁源生产部-餐厨厂运行班,珠海市海宜洁源餐厨垃圾处置有限公司-洁源生产部-餐厨厂地磅',['部门']]='珠海市海宜洁源餐厨垃圾处置有限公司-洁源生产部-餐厨厂运行班'
hebing.loc[hebing['部门']=='珠海市海宜洁源餐厨垃圾处置有限公司-洁源生产部-餐厨厂机修班',['部门']]='珠海市海宜洁源餐厨垃圾处置有限公司-洁源生产部-洁源机修班'
hebing.loc[hebing['部门']=='珠海市海宜洁源餐厨垃圾处置有限公司-洁源生产部-餐厨厂运行班',['部门']]='珠海市海宜洁源餐厨垃圾处置有限公司-洁源生产部-洁源运行班'
hebing.loc[hebing['部门']=='珠海市海宜洁源餐厨垃圾处置有限公司-洁源生产部',['部门']]='珠海市海宜洁源餐厨垃圾处置有限公司-洁源生产部-洁源生产部'
hebing.loc[hebing['部门']=='珠海市海宜洁源餐厨垃圾处置有限公司',['部门']]='珠海市海宜洁源餐厨垃圾处置有限公司-洁源总经理-洁源总经理'

#姓名替换
hebing.loc[hebing['学员姓名']=='周家魁',['学员姓名']]='周贾魁' # 用于排序1，2字相同的姓名
hebing.loc[hebing['学员姓名']=='梁敏豪',['学员姓名']]='梁民豪' # 用于排序1，2字相同的姓名
hebing.loc[hebing['学员姓名']=='赵文辉',['学员姓名']]='赵温辉' # 用于排序1，2字相同的姓名
hebing.loc[hebing['学员姓名']=='邱德怀',['学员姓名']]='邱嘚怀' # 用于排序1，2字相同的姓名
hebing.loc[hebing['学员姓名']=='赵文富',['学员姓名']]='赵问富' # 用于排序1，2字相同的姓名

#职位替换
hebing.loc[hebing['职位']=='收运专责',['职位']]='专责'
hebing.loc[hebing['职位']=='地磅员',['职位']]='运行工'
hebing.loc[hebing['职位'].isna(),['职位']]='测试盒子'         # 如果为nan，则替换该格子为‘测试格子’

#分辨收运部的东西区
hebing.loc[(hebing['职位']=='司机') & (hebing['部门']=='珠海市海宜洁源餐厨垃圾处置有限公司-洁源收运部-洁源东区服务点'),['部门']]='珠海市海宜洁源餐厨垃圾处置有限公司-洁源收运部-洁源东司机'
hebing.loc[(hebing['职位']=='司机') & (hebing['部门']=='珠海市海宜洁源餐厨垃圾处置有限公司-洁源收运部-洁源东区服务点-东区司机'),['部门']]='珠海市海宜洁源餐厨垃圾处置有限公司-洁源收运部-洁源东司机'
hebing.loc[(hebing['职位']=='司机') & (hebing['部门']=='珠海市海宜洁源餐厨垃圾处置有限公司-洁源收运部-洁源西区服务点'),['部门']]='珠海市海宜洁源餐厨垃圾处置有限公司-洁源收运部-洁源西司机'
hebing.loc[(hebing['职位']=='司机') & (hebing['部门']=='珠海市海宜洁源餐厨垃圾处置有限公司-洁源收运部-洁源西区服务点-西区司机,珠海市海宜洁源餐厨垃圾处置有限公司-洁源收运部-洁源公司合同'),['部门']]='珠海市海宜洁源餐厨垃圾处置有限公司-洁源收运部-洁源西司机'
hebing.loc[(hebing['职位']=='司机') & (hebing['部门']=='珠海市海宜洁源餐厨垃圾处置有限公司-洁源收运部-洁源公司合同,珠海市海宜洁源餐厨垃圾处置有限公司-洁源收运部-洁源东区服务点-东区司机'),['部门']]='珠海市海宜洁源餐厨垃圾处置有限公司-洁源收运部-洁源东司机'

hebing.loc[(hebing['职位']=='押运员') & (hebing['部门']=='珠海市海宜洁源餐厨垃圾处置有限公司-洁源收运部-洁源公司押运,珠海市海宜洁源餐厨垃圾处置有限公司-洁源收运部-洁源东区服务点-东区押运员'),['部门']]='珠海市海宜洁源餐厨垃圾处置有限公司-洁源收运部-洁源动押运员'
hebing.loc[(hebing['职位']=='押运员') & (hebing['部门']=='珠海市海宜洁源餐厨垃圾处置有限公司-洁源收运部-洁源西区服务点-西区押运员,珠海市海宜洁源餐厨垃圾处置有限公司-洁源收运部-洁源公司押运'),['部门']]='珠海市海宜洁源餐厨垃圾处置有限公司-洁源收运部-洁源喜押运员'
hebing.loc[(hebing['职位']=='押运员') & (hebing['部门']=='珠海市海宜洁源餐厨垃圾处置有限公司-洁源收运部-洁源西区服务点-西区押运员,珠海市海宜洁源餐厨垃圾处置有限公司-洁源收运部-洁源公司合同'),['部门']]='珠海市海宜洁源餐厨垃圾处置有限公司-洁源收运部-洁源喜押运员'

hebing.loc[(hebing['职位']=='专责') & (hebing['部门']=='珠海市海宜洁源餐厨垃圾处置有限公司-洁源收运部-洁源西区服务点-西区押运员,珠海市海宜洁源餐厨垃圾处置有限公司-洁源收运部-洁源西区服务点-西区司机,珠海市海宜洁源餐厨垃圾处置有限公司-洁源收运部-洁源西区服务点'),['部门']]='珠海市海宜洁源餐厨垃圾处置有限公司-洁源收运部-洁源喜押运员'
hebing.loc[(hebing['职位']=='专责') & (hebing['部门']=='珠海市海宜洁源餐厨垃圾处置有限公司-洁源收运部-洁源公司押运,珠海市海宜洁源餐厨垃圾处置有限公司-洁源收运部-洁源公司合同,珠海市海宜洁源餐厨垃圾处置有限公司-洁源收运部-洁源东区服务点-东区押运员'),['部门']]='珠海市海宜洁源餐厨垃圾处置有限公司-洁源收运部-洁源东押运员'
hebing.loc[(hebing['职位']=='分公司部门经理') & (hebing['部门']=='珠海市海宜洁源餐厨垃圾处置有限公司-洁源收运部-洁源公司押运,珠海市海宜洁源餐厨垃圾处置有限公司-洁源收运部-洁源公司合同,珠海市海宜洁源餐厨垃圾处置有限公司-洁源收运部'),['部门']]='珠海市海宜洁源餐厨垃圾处置有限公司-洁源收运部-洁源东押运员'

# ============排序，按照拼音升序================================================
# 增加多余的列来放姓名的首字拼音字母，排序后，将该列删除
hebing['a'] = hebing['部门'].apply(lambda x: lazy_pinyin(x)[20])
hebing['b'] = hebing['部门'].apply(lambda x: lazy_pinyin(x,style=pypinyin.TONE2)[26])
#hebing['c'] = hebing['职位'].apply(lambda x: lazy_pinyin(x)[0])
hebing['x'] = hebing['学员姓名'].apply(lambda x: lazy_pinyin(x,style=pypinyin.TONE2)[0])  # lazy_pinyin(x)[0][1]：第0字，字的第1个字母，删除标识全要；调用拼音库
hebing['y'] = hebing['学员姓名'].apply(lambda x: lazy_pinyin(x,style=pypinyin.TONE2)[1])  # 增加多余的列用来标识姓名的字母，使用声调转数字以加强代码
                                                                                          # apply:沿dataframe的轴应用函数（沿dataframe按每一行或列应用函数）

#hebing = hebing.sort_values(by=['c'], ascending=True)
hebing = hebing.sort_values(by=['指派/加入时间'], ascending=True)
hebing = hebing.sort_values(by=['x','y'       ], ascending=True)
hebing = hebing.sort_values(by=['a','b'       ], ascending=True)
  
# dataframe排序
hebing = hebing.reset_index(drop=True) # 重置索引，避免对后面造成混乱

# 删除多余列
hebing.drop(columns='x', inplace=True) # 删除多余的列
hebing.drop(columns='y', inplace=True)
hebing.drop(columns='a', inplace=True)
hebing.drop(columns='b', inplace=True)
# hebing.drop(columns='c', inplace=True)

# =============按完成情况分类，按时完成、未开始、进行中============================
###########按时完成############
anshiwancheng = hebing[hebing['完成情况'] == '按时完成']      # 按条件筛选
#anshiwancheng = anshiwancheng.dropna()                      # 删除weiaishi dataframe 中的nan
anshiwancheng = anshiwancheng.reset_index(drop=True)         # 重置索引，避免对后面造成混乱
############未开始############
weikaishi     = hebing[hebing['完成情况'] == '未开始']            
weikaishi     = weikaishi.reset_index(drop=True)
############进行中############
jinxingzhong  = hebing[hebing['完成情况'] == '进行中']            
jinxingzhong  = jinxingzhong.reset_index(drop=True)

# ================储存为d月份的培训教育excel=====================================
with pd.ExcelWriter(f'【{d}】月洁源公司安全教育培训.xlsx') as writer:
    anshiwancheng.to_excel(writer, sheet_name='按时完成')
    weikaishi.    to_excel(writer, sheet_name='未开始')
    jinxingzhong. to_excel(writer, sheet_name='进行中')
    hebing.       to_excel(writer, sheet_name='总数据')
print(os.listdir())

# =============================================================================
# ================Dataframe 转换为word=========================================
# =============================================================================

'''字体大小对应pt

初号44pt
小初36pt
号26pt
小一24pt
二号22pt
小二18pt
三号16pt
小三15pt
四号14pt
小四12pt
五号10.5pt
八号5
七号5.5
字号‘小六’对应磅值6.5
字号‘六号’对应磅值7.5
字号‘小五’对应磅值9
'''
# =============================================================================
# 直接读取docx模板填写数据，避免后期表格有变动后需要大改
# =============================================================================

#在按时完成的dataframe内最后行增加一行无效行，为了能完全输出最后一行，同时不影响excel输出及下方单页生成姓名。
print(f'{anshiwancheng}共{len(anshiwancheng)}行')

anshiwancheng.loc[len(anshiwancheng)] = [*["此页面不打印"]*(len(anshiwancheng.columns)-1),"0"]#放0方便成绩行转换，.astype(int)不能转换nan

#由于python-docx无 copy table 功能的函数，因此直接在外部word内复制好表格
peixunjilu = Document(r"E:\海宜洁源安全共享文件\标准化自评\【5】教育培训\5.1教育培训管理\3-培训记录及评估\1-日常培训记录\一人一档内容编制\0-模板\一人一档月度模板2.docx") #word模板路径
peixunjilu.tables
t =0
for table in peixunjilu.tables: 
    t = t + 1
print('表格数量:')
print(t)

#将dataframe转换成list
riqi     = anshiwancheng['指派/加入时间'].values.tolist()
kecheng  = anshiwancheng['课程名称'     ].values.tolist() 
keshi    = anshiwancheng['课时'         ].values.tolist()
gangwei  = anshiwancheng['职位'         ].values.tolist() 
xingming = anshiwancheng['学员姓名'     ].values.tolist() 
bumen    = anshiwancheng['部门'         ].values.tolist() 
cheng    = anshiwancheng['最新一次成绩' ].fillna(0)     # 将dataframe中的nan替换成0，因为.astype(int)不能转换nan
chengj   = cheng.astype(int).tolist()                   # 先转换为整型，再转列表
tj       = len(chengj)
chengjii = list(map(str, chengj))                       # 将list转字符串
chengji  = [' ' if i =='0' else i for i in chengjii]    # 将list中的字符串“nan”替换成空格

a1 = 1 # cell控制参数
c1 = 0 # tables控制参数
p1 = 0 # paragraphs控制参数
f1 = 0 # format内的函数控制参数


print('按时完成的行数：')
print(len(anshiwancheng))
for i in range(len(anshiwancheng)): # range(10) 是 0 到 9；len 统计dataframe的index数量，行数
    if i  == len(anshiwancheng)-1:  # 跳出循环
        break
    
    #培训日期
    B_riqi = peixunjilu.tables[c1].cell(a1,0).paragraphs[0].add_run('{}'.format(riqi[i])[0:10])#切片字符串【0：10】
    B_riqi.font.name = '宋体'  ##
    B_riqi.font.size = Pt(12)  ##
    r3 = B_riqi._element 
    r3.rPr.rFonts.set(qn('w:eastAsia'), '宋体') ##中文字体

    #培训课程
    B_riqi = peixunjilu.tables[c1].cell(a1,1).paragraphs[0].add_run('{}'.format(kecheng[i][1:-12]))
    B_riqi.font.name = '宋体'  ##
    B_riqi.font.size = Pt(12)  ##
    r3 = B_riqi._element 
    r3.rPr.rFonts.set(qn('w:eastAsia'), '宋体') ##中文字体

    #课时
    B_riqi = peixunjilu.tables[c1].cell(a1,2).paragraphs[0].add_run('{}'.format(keshi[i]))
    B_riqi.font.name = '宋体'  ##
    B_riqi.font.size = Pt(12)  ##
    r3 = B_riqi._element 
    r3.rPr.rFonts.set(qn('w:eastAsia'), '宋体') ##中文字体

    #授课人
    B_riqi = peixunjilu.tables[c1].cell(a1,3).paragraphs[0].add_run('钉钉线上授课')
    B_riqi.font.name = '宋体'  ##
    B_riqi.font.size = Pt(12)  ##
    r3 = B_riqi._element 
    r3.rPr.rFonts.set(qn('w:eastAsia'), '宋体') ##中文字体

    #受教育人岗位
    B_riqi = peixunjilu.tables[c1].cell(a1,4).paragraphs[0].add_run('{}'.format(gangwei[i]))
    B_riqi.font.name = '宋体'  ##
    B_riqi.font.size = Pt(12)  ##
    r3 = B_riqi._element 
    r3.rPr.rFonts.set(qn('w:eastAsia'), '宋体') ##中文字体
    
    #考试成绩
    B_riqi = peixunjilu.tables[c1].cell(a1,5).paragraphs[0].add_run('{}'.format(chengji[i]))
    B_riqi.font.name = '宋体'  ##
    B_riqi.font.size = Pt(12)  ##
    r3 = B_riqi._element 
    r3.rPr.rFonts.set(qn('w:eastAsia'), '宋体') ##中文字体
    
    f1 = f1 + 1
    
    if xingming[i] == xingming[i+1]:
        # print('该行与上一行相等')
        
        a1 = a1 + 1 # cell控制参数,行 ：相等后，行+1
        # 当相等时，不写paragraphs
        
    else:# 如果名字不相等，在下一页添加内容
         # print('该行与上一行不等')
        a1 = 1     # 不相等后，将a1复原
        c1 = c1 + 1# tables控制参数
        
        #页面底部的  姓名 部门等文本框
        B_riqi = peixunjilu.paragraphs[p1].add_run('{}'.format(xingming[i]))
        B_riqi.font.name = '宋体' ##
        B_riqi.font.size = Pt(6)  ##
        r3 = B_riqi._element 
        r3.rPr.rFonts.set(qn('w:eastAsia'), '宋体') ##中文字体
        B_riqi.font.color.rgb = RGBColor(200, 200, 200)

        B_riqi = peixunjilu.paragraphs[p1].add_run('{}'.format(bumen[i][17:]))
        B_riqi.font.name = '宋体'  ##
        B_riqi.font.size = Pt(6)  ##
        r3 = B_riqi._element 
        r3.rPr.rFonts.set(qn('w:eastAsia'), '宋体') ##中文字体
        B_riqi.font.color.rgb = RGBColor(200, 200, 200)
        
        p1 = p1 + 2 # paragraphs的控制参数
        
peixunjilu.save('洁源公司{}月一人一档.docx'.format(d))# 保存
#
# =======================文本框修改===============================================
# document = Document('11.docx')
# children = document.element.body.iter()
# for child in children:
#     if child.tag.endswith('txbx'):            # child.tag.endswith('textbox') or txbxContent
#         for ci in child.iter():
#             if ci.tag.endswith('main}t'):
#                 ci.text = '我被修改了！！'
#  document.save('111.docx')
# ========================文本框失效=============================================
# 创建文档
#peixunjilu = Document()


#修改页边距changing the page margins
#sections = peixunjilu.sections 
#for section in sections: 
#    section.top_margin = Cm(1.5) 
#    section.bottom_margin = Cm(0.5) 
#    section.left_margin = Cm(2) 
#    section.right_margin = Cm(2.7) ##Cm 需要import 一下

# 标题
#peixunjilu.add_heading('标题0', 0)
#peixunjilu.add_heading('标题1', 1)
#peixunjilu.add_heading('标题2', 2)

#页眉

#header = peixunjilu.sections[0].header # 获取第一个节的页眉（所有的页眉都一致）
#peixunjilu.sections[0].even_page_header.paragraphs[0].text = "这是偶数页页眉"
#peixunjilu.sections[0].header.paragraphs[0].text = "这是奇数页页眉"


##在页眉中插入表格
#table_h = header.add_table(rows=3, cols=2,width=Cm(15.95))##3行2列的表格 
#table_h.alignment = WD_TABLE_ALIGNMENT.CENTER#将表格居中

#header.paragraphs[0].paragraph_format.space_after = Pt(0)#段后行距
#header.paragraphs[0].clear()


#cell_yemei_1 = table_h.cell(0,0).merge(table_h.cell(1,0))#合并单元格
#cell_yemei_2 = table_h.cell(2,0).merge(table_h.cell(2,1))
#cell_yemei_1.width = Cm(12.7)
#table_h.cell(0,1).width = Cm(3.25)#单元格宽度
#table_h.cell(1,1).width = Cm(3.25)
#设置单元格宽度及行高
#table_h.rows[0].height= Cm(0.5)#单元格高度
# table_h.rows[1].height= Cm(0.5)
# table_h.rows[2].height= Cm(0.99)
# 
# table_h.cell(0,0).paragraphs[0].paragraph_format.space_after = Pt(0)#段落后间距
# table_h.cell(0,1).paragraphs[0].paragraph_format.space_after = Pt(0)
# table_h.cell(1,1).paragraphs[0].paragraph_format.space_after = Pt(0)
# table_h.cell(2,0).paragraphs[0].paragraph_format.space_after = Pt(0)
# 
# table_h.cell(0,0).paragraphs[0].paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE#单倍行距
# table_h.cell(0,1).paragraphs[0].paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE#单倍行距
# table_h.cell(1,1).paragraphs[0].paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE#单倍行距
# table_h.cell(2,0).paragraphs[0].paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE#单倍行距




# ========================段落设置================================================
# paragraph1 = doc.add_paragraph("这是一个普通的段落这是一个普通的段落这是一个普通的段落")
# # paragraph.line_spacing_rule = WD_LINE_SPACING.EXACTLY  # 行距固定值
# # paragraph.line_spacing_rule = WD_LINE_SPACING.MULTIPLE  # 多倍行距
# # paragraph1.paragraph_format.line_spacing = 1.5  # 行间距，1.5倍行距
# paragraph1.paragraph_format.line_spacing = Pt(20)  # 行间距，固定值20磅
# paragraph1.paragraph_format.first_line_indent = Pt(10)  # 首行缩进10磅
# paragraph1.paragraph_format.space_before = Pt(30)  # 段前30磅
# paragraph1.paragraph_format.space_after = Pt(15)  # 段后15磅

##小白学python—处理word中的段落https://zhuanlan.zhihu.com/p/372902503
# =============================================================================

##在表格中填入文字,run是指段落中某一部分文本，我们获取到某部分文本才能设置其样式
#run = table_h.cell(0,0).paragraphs[0].add_run(u'珠海市海宜环境投资有限公司安全生产管理体系')
#run.font.size = Pt(12)
#run.italic = False ##设置是否为斜体
#run.bold = False#加粗
#run.underline = False#下划线
# run.font.underline = WD_UNDERLINE.DOUBLE  # 设置为双下划线
# 查看所有下划线类型
# for line_type in WD_UNDERLINE.__members__:
#     print(line_type.name)
#run1.font.shadow = True  # 是否阴影
#run.font.strike = True  # 是否删除线
# run.font.double_strike = True  # 是否双删除线
#run1.font.color.rgb = RGBColor(56, 36, 255)  # 字体颜色
#run.font.highlight_color = WD_COLOR_INDEX.YELLOW  # 文本高亮颜色，此次设置为黄色
# 查看所有支持的高亮颜色
# for color in WD_COLOR_INDEX.__members__:
#     print(color.name)



# run.font.name ='黑体'# 设置西文是新罗马字体"Times New Roman"
# r = run._element#需用该方法做中介才能设置字体
# r.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
# table_h.cell(0,0).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER#页眉内表格（0，0）内容居中
# 
# ##第一个cell中的文字
# run1 = table_h.cell(0,1).paragraphs[0].add_run(u'编号：HY-PX006')
# run1.font.name = '宋体' ##
# run1.font.size = Pt(9)  ##
# run1.bold = False  ##加粗
# r1 = run1._element 
# r1.rPr.rFonts.set(qn('w:eastAsia'), '宋体') ##中文字体
# table_h.cell(0,1).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY#两端对齐
# ##第二个cell中的文字
# run2 = table_h.cell(1,1).paragraphs[0].add_run(u'版本/修订：A/1')
# run2.font.name = '宋体' ##
# run2.font.size = Pt(9)  ##
# run2.bold = False  ##加粗
# r2 = run2._element 
# r2.rPr.rFonts.set(qn('w:eastAsia'), '宋体') ##中文字体
# table_h.cell(1,1).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY#两端对齐
# 
# run3 = table_h.cell(2,0).paragraphs[0].add_run(u'违规、违章操作记录')
# run3.font.name = '宋体' ##
# run3.font.size = Pt(18)  ##
# run3.bold = True  ##加粗
# r3 = run3._element 
# r3.rPr.rFonts.set(qn('w:eastAsia'), '宋体') ##中文字体
# table_h.cell(2,0).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER#两端对齐
# 
# 
# peixunjilu.sections[0].header_distance = Cm(1.5) #页眉距顶端距离
# peixunjilu.sections[0].footer_distance = Cm(0.8)#页眉距低端距离
# 
# def remove_row(table, row):#删除表格中某一行
#     tbl = table._tbl
#     tr = row._tr
#     tbl.remove(tr)










